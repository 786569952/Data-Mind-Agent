"""
data_platform/document.py
=========================
非结构化文档处理 (模拟数据平台 ODS -> ETL 中的文档流支线)。

支持文件类型:
    - .docx  (Word 文档，python-docx 解析)
    - .md    (Markdown / PRD，按标题切分 + 递归字符切分)
    - .txt   (纯文本)
    - .prd   (产品需求文档，按 Markdown 规范解析)

流程:
    1. parse_document()   : 文件 -> 纯文本 (+ 元信息)
    2. split_text()       : 纯文本 -> 切分块 (chunk)，参数可调
    3. embed_chunks()     : 切分块 -> 向量，逐步产生 EmbeddingStep 日志
"""
from __future__ import annotations

import os
import shutil
from dataclasses import dataclass, field
from typing import Any, Iterator

from langchain_core.documents import Document
from langchain_core.embeddings import Embeddings
from langchain_text_splitters import (
    RecursiveCharacterTextSplitter,
    MarkdownHeaderTextSplitter,
)


# ----------------------------- 数据结构 -----------------------------
@dataclass
class ParseInfo:
    """文档解析元信息。"""
    filename: str
    file_type: str           # docx / md / txt / prd
    total_chars: int
    n_paragraphs: int
    n_tables: int = 0        # docx 表格数
    headers: list[str] = field(default_factory=list)  # md/prd 标题结构


@dataclass
class ChunkInfo:
    """单个切分块信息。"""
    index: int
    text: str
    char_len: int
    metadata: dict[str, Any] = field(default_factory=dict)


@dataclass
class EmbeddingStep:
    """单条 chunk 的向量化记录，供前端逐步展示。"""
    index: int
    text_preview: str        # 前 80 字
    char_len: int
    vector_dim: int          # 向量维度
    vector_preview: list[float]  # 前 5 个分量
    status: str = "ok"       # ok / error
    error: str = ""


@dataclass
class EmbeddingReport:
    """整轮 embedding 的汇总报告。"""
    model: str
    total_chunks: int
    success: int
    failed: int
    vector_dim: int
    steps: list[EmbeddingStep] = field(default_factory=list)
    total_tokens: int = 0    # 估算


# ----------------------------- 1. 文档解析 -----------------------------
def parse_document(file_path: str, filename: str = "") -> tuple[str, ParseInfo]:
    """
    解析文档为纯文本，返回 (文本, 解析信息)。

    支持 doc / docx / md / txt / prd。
    .doc (旧版 Word) 通过 macOS textutil 转换为 txt 后解析。
    """
    filename = filename or os.path.basename(file_path)
    ext = os.path.splitext(filename)[1].lower().lstrip(".")
    if ext not in ("doc", "docx", "md", "txt", "prd"):
        raise ValueError(f"暂不支持的文档类型: .{ext}")

    if ext == "doc":
        # 旧版 Word .doc: 用 macOS 自带 textutil 转成 txt
        import subprocess
        import tempfile
        if not shutil.which("textutil"):
            raise ValueError(
                "解析 .doc 需要 macOS textutil 工具。"
                "请将文件另存为 .docx 格式后重试。"
            )
        with tempfile.NamedTemporaryFile(
            suffix=".txt", delete=False, mode="w"
        ) as tmp:
            txt_path = tmp.name
        try:
            subprocess.run(
                ["textutil", "-convert", "txt", "-encoding", "utf-8",
                 file_path, "-output", txt_path],
                check=True, capture_output=True,
            )
            with open(txt_path, "r", encoding="utf-8") as f:
                text = f.read()
            info = ParseInfo(
                filename=filename, file_type="doc",
                total_chars=len(text),
                n_paragraphs=len([p for p in text.split("\n\n") if p.strip()]),
            )
            return text, info
        finally:
            try:
                os.unlink(txt_path)
            except Exception:
                pass

    if ext == "docx":
        text, info = _parse_docx(file_path, filename)
    elif ext in ("md", "prd"):
        text, info = _parse_markdown(file_path, filename, ext)
    else:  # txt
        with open(file_path, "r", encoding="utf-8") as f:
            text = f.read()
        info = ParseInfo(
            filename=filename, file_type="txt",
            total_chars=len(text),
            n_paragraphs=len([p for p in text.split("\n\n") if p.strip()]),
        )

    return text, info


def _parse_docx(path: str, filename: str) -> tuple[str, ParseInfo]:
    import docx

    doc = docx.Document(path)
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    text = "\n\n".join(paragraphs)

    # 提取表格
    table_texts: list[str] = []
    for tbl in doc.tables:
        rows = []
        for row in tbl.rows:
            cells = [c.text.strip() for c in row.cells]
            rows.append(" | ".join(cells))
        if rows:
            table_texts.append("\n".join(rows))
    if table_texts:
        text += "\n\n" + "\n\n".join(table_texts)

    info = ParseInfo(
        filename=filename, file_type="docx",
        total_chars=len(text),
        n_paragraphs=len(paragraphs),
        n_tables=len(doc.tables),
    )
    return text, info


def _parse_markdown(path: str, filename: str, ext: str) -> tuple[str, ParseInfo]:
    with open(path, "r", encoding="utf-8") as f:
        text = f.read()

    # 提取标题结构
    headers = [
        line.strip() for line in text.splitlines()
        if line.strip().startswith("#")
    ]
    info = ParseInfo(
        filename=filename, file_type=ext,
        total_chars=len(text),
        n_paragraphs=len([p for p in text.split("\n\n") if p.strip()]),
        headers=headers,
    )
    return text, info


# ----------------------------- 2. 切分 -----------------------------
def split_text(
    text: str,
    chunk_size: int = 500,
    chunk_overlap: int = 50,
    separators: list[str] | None = None,
    is_markdown: bool = False,
    source_filename: str = "",
) -> list[ChunkInfo]:
    """
    按可调参数切分文本，返回 ChunkInfo 列表。

    - is_markdown=True 时，先按标题切分，再按字符切分，保留标题层级元数据。
    - separators 默认: ["\n\n", "\n", "。", "！", "？", "；", ".", "!", "?", ";", " ", ""]
    """
    if separators is None:
        separators = ["\n\n", "\n", "。", "！", "？", "；", ".", "!", "?", ";", " ", ""]

    if is_markdown:
        # 先按 Markdown 标题切，得到带标题元数据的片段
        md_splitter = MarkdownHeaderTextSplitter(
            headers_to_split_on=[
                ("#", "h1"), ("##", "h2"), ("###", "h3"),
            ],
            strip_headers=False,
        )
        try:
            md_docs = md_splitter.split_text(text)
        except Exception:
            md_docs = [Document(page_content=text)]

        # 再对每个片段做字符级递归切分
        char_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            separators=separators,
            keep_separator=True,
        )
        chunks: list[ChunkInfo] = []
        idx = 0
        for md_doc in md_docs:
            sub_texts = char_splitter.split_text(md_doc.page_content)
            for st in sub_texts:
                if not st.strip():
                    continue
                chunks.append(ChunkInfo(
                    index=idx,
                    text=st,
                    char_len=len(st),
                    metadata={
                        "source": source_filename,
                        **md_doc.metadata,
                    },
                ))
                idx += 1
        return chunks

    # 普通文本递归切分
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        separators=separators,
        keep_separator=True,
    )
    texts = splitter.split_text(text)
    return [
        ChunkInfo(
            index=i, text=t, char_len=len(t),
            metadata={"source": source_filename},
        )
        for i, t in enumerate(texts) if t.strip()
    ]


# ----------------------------- 3. 向量化 (带过程追踪) -----------------------------
def embed_chunks(
    chunks: list[ChunkInfo],
    embeddings: Embeddings,
    model_name: str = "",
    batch_size: int = 8,
) -> tuple[list[Document], EmbeddingReport]:
    """
    逐步向量化 chunks，返回 (LangChain Documents, EmbeddingReport)。

    每批调用 embeddings.embed_documents，并产出 EmbeddingStep 日志，
    前端可据此展示 "embedding 全过程"。
    """
    report = EmbeddingReport(
        model=model_name,
        total_chunks=len(chunks),
        success=0,
        failed=0,
        vector_dim=0,
    )

    if not chunks:
        return [], report

    all_vecs: list[list[float]] = []
    for start in range(0, len(chunks), batch_size):
        batch = chunks[start:start + batch_size]
        batch_texts = [c.text for c in batch]
        try:
            vecs = embeddings.embed_documents(batch_texts)
        except Exception as e:
            # 整批失败，逐条降级
            for c in batch:
                try:
                    v = embeddings.embed_documents([c.text])[0]
                    all_vecs.append(v)
                    report.steps.append(_make_step(c, v, "ok"))
                    report.success += 1
                    report.vector_dim = len(v)
                except Exception as e2:
                    all_vecs.append([0.0])
                    report.steps.append(_make_step(c, [], "error", str(e2)))
                    report.failed += 1
            continue

        for c, v in zip(batch, vecs):
            all_vecs.append(v)
            report.steps.append(_make_step(c, v, "ok"))
            report.success += 1
            report.vector_dim = len(v)

    # 估算 token (中文按 1.5 字/token 粗估)
    report.total_tokens = int(sum(c.char_len for c in chunks) / 1.5)

    # 构造 LangChain Document
    docs = [
        Document(page_content=c.text, metadata={
            "chunk_index": c.index,
            "char_len": c.char_len,
            **c.metadata,
        })
        for c in chunks
    ]
    return docs, report


def _make_step(chunk: ChunkInfo, vec: list[float],
               status: str, error: str = "") -> EmbeddingStep:
    return EmbeddingStep(
        index=chunk.index,
        text_preview=chunk.text[:80].replace("\n", " "),
        char_len=chunk.char_len,
        vector_dim=len(vec),
        vector_preview=vec[:5] if vec else [],
        status=status,
        error=error,
    )
